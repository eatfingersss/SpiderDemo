import os
import re
import docx
import GetPs

dir = os.listdir('./')
option = ('docx','doc')
matchobj1 = '([0-9]+)*.* *'
matchobj2 = '([a-z]|[A-Z])[a-z]+(\t| )+'
matchobj2_ = '\[.*\]'
matchobj3 = '[a-z]\.?'

matchobj = (matchobj1 + matchobj2 + matchobj2_ + matchobj3, matchobj1 + matchobj2 + matchobj3)

match_font = (matchobj1 + matchobj2 + matchobj2_, matchobj1 + matchobj2)
match_rear = (matchobj3,)
# 临时加的
match_font = match_font[0]


def is_it_that_name(name):
    for one in option:
        if one == name:
            return True
    return False


def is_it_that_paragraphs(str):
    if re.match(matchobj[0], str):
        return False  # 已有音标
    elif re.match(matchobj[1], str):
        # 找出位置
        return re.match('.*' + matchobj2, str).regs[0]
        # 没有音标
    else:
        return False  # 不规范


for one in dir:
    try:
        filenames = one.split('.')
        if filenames.__len__() is not 2:
            raise NameError('有多个.')
        if is_it_that_name(filenames[1]) is False:
            raise NameError('不是要找的文件')
    except Exception as e:
        continue

    try:
        # 获取文档对象
        file = docx.Document(one)
    except Exception as e:
        print("文件" + one + "打开失败...")
        continue
    print("文件 " + one + ":")
    # 输出每一段的内容
    i = 1
    for para in file.paragraphs:
        result = is_it_that_paragraphs(para.text)
        if result is not False:
            # 操作
            try:
                # print("第" + str(i) + "段的内容是：" + file.paragraphs[i - 1].text)
                # 将前段与后端分割
                words = str(para.text[result[0]  # 0  # re.match('^([0-9]+)*.* *', para.text).regs[0][1]
                                      :
                                      result[1]  # re.match(match_font, para.text).regs[0][1]
                            ]
                            )
                if words == '':
                    raise Exception('单词提炼失败')
                words = words.strip()
                # tail = re.match(match_font + ' +', para.text).regs[0][1]
                if words == '':
                    raise Exception('单词筛选失败')
                print('\t', end='')
                ps = GetPs.main(words)
                if ps == '' or ps is None:
                    raise Exception('音标获取失败')
                new_para = str(para.text[:result[1]]).rstrip() + '\t' + ps + para.text[result[1]:]
                p = file.paragraphs[i - 1].clear()
                file.paragraphs[i - 1].add_run(new_para)

                print("  已添加")
            except Exception as e:
                print("\n[error]单词 " + words + " 所在列添加失败:" + e.__str__())

        i += 1
    file.save(one)
    print(one + " 已保存，共扫描了"+str(i)+"行\n")
os.system('pause')